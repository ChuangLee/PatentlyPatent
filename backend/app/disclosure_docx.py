"""按 No.34 模板生成专利技术交底书 .docx（python-docx 实现）

输入：
  - project: ORM Project 对象
  - all_md_files: list[FileNode]，已生成的 markdown 文件
        命名规范（与 mining.py 对齐）：
            01-背景技术.md
            02-现有技术缺点.md
            03-技术问题.md
            04-技术方案.md
            05-关键点.md
            06-优点.md
            _问题清单.md

输出：bytes（docx 二进制），由调用方决定持久化位置 / FileNode 落库。

markdown → docx 段落映射（自手写解析，零三方依赖）：
  - 「# 」→ heading 1，「## 」→ heading 2，「### 」→ heading 3
  - 「- 」「* 」「• 」→ list paragraph（无嵌套）
  - 「| col | col |」表格行 → 拼为纯文本段（python-docx 表格构造较重，先 stringify）
  - 其他 → 普通段落，空行作为段落分隔
  - 行内 **bold** 简单 run 标粗
"""
from __future__ import annotations

import io
import re
from datetime import datetime
from typing import Iterable, List, Optional

from docx import Document
from docx.document import Document as _Doc
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


# ─────────────────────────── 常量 ───────────────────────────

DOMAIN_LABEL = {
    "cryptography": "密码学",
    "infosec": "信息安全",
    "ai": "人工智能",
    "other": "其他",
}

PLACEHOLDER = "（待 AI 进一步引导填写）"

# 文件名 → 章节 key 映射；按优先级匹配
SECTION_FILE_KEYWORDS = {
    "background":   ["背景技术", "01-"],
    "disadvantage": ["现有技术缺点", "缺点", "02-"],
    "problem":      ["技术问题", "03-"],
    "solution":     ["技术方案", "04-"],
    "key_points":   ["关键点", "05-"],
    "advantage":    ["优点", "06-"],
    "questions":    ["问题清单", "_问题"],
}


# ─────────────────────────── 工具 ───────────────────────────

def _classify_md(name: str) -> Optional[str]:
    """根据文件名归类到章节 key。"""
    for key, kws in SECTION_FILE_KEYWORDS.items():
        for kw in kws:
            if kw in name:
                return key
    return None


def _strip_inline_md(text: str) -> str:
    """去除最常见的行内 md 标记，保留可读文本（粗体/斜体/code/链接）。"""
    # ![alt](url) → alt
    text = re.sub(r"!\[([^\]]*)\]\([^)]+\)", r"\1", text)
    # [text](url) → text (url)
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", text)
    # `code`
    text = re.sub(r"`([^`]+)`", r"\1", text)
    # **bold** / __bold__
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"__([^_]+)__", r"\1", text)
    # *italic* / _italic_  （避免误伤数学下划线，做保守清理）
    text = re.sub(r"(?<!\*)\*([^*\n]+)\*(?!\*)", r"\1", text)
    return text.strip()


def _add_para(doc: _Doc, text: str, *, size: int = 11, bold: bool = False,
              color: Optional[str] = None, italic: bool = False,
              align: Optional[int] = None, after: int = 4) -> None:
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    run = p.add_run(text or "")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def _add_heading(doc: _Doc, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        # heading 默认字号在中文环境下偏小，简单放大
        if level == 1:
            r.font.size = Pt(16)
        elif level == 2:
            r.font.size = Pt(14)
        else:
            r.font.size = Pt(12)


def _add_meta(doc: _Doc, text: str) -> None:
    """灰色斜体「提示」段。"""
    _add_para(doc, text, size=9, color="888888", italic=True, after=2)


def _md_to_paragraphs(doc: _Doc, md: str, *, base_heading_offset: int = 1) -> None:
    """把一段 markdown 文本附加到 doc。

    base_heading_offset：md 中的 「# 」对应到 docx 的哪一级 heading。
    传 1 时：md 「# X」→ heading 1；md 「## X」→ heading 2 …
    传 2 时：md 「# X」→ heading 2 …
    """
    if not md or not md.strip():
        return

    lines = md.splitlines()
    in_table = False
    table_buf: List[str] = []

    def _flush_table() -> None:
        nonlocal in_table, table_buf
        if not table_buf:
            return
        # 简化处理：把每一行当成一段，列以 ｜ 分隔保留可读
        for row in table_buf:
            cells = [c.strip() for c in row.strip().strip("|").split("|")]
            # 跳过分隔行 |---|---|
            if all(re.fullmatch(r":?-+:?", c) for c in cells if c):
                continue
            _add_para(doc, "  ｜  ".join(cells), size=10, after=2)
        table_buf = []
        in_table = False

    i = 0
    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()

        # 表格行
        if line.lstrip().startswith("|") and line.rstrip().endswith("|"):
            in_table = True
            table_buf.append(line)
            i += 1
            continue
        elif in_table:
            _flush_table()

        # 空行
        if not line.strip():
            i += 1
            continue

        # heading
        m = re.match(r"^(#{1,6})\s+(.+)$", line)
        if m:
            depth = len(m.group(1))
            text = _strip_inline_md(m.group(2))
            level = min(9, base_heading_offset + depth - 1)
            _add_heading(doc, text, level=level)
            i += 1
            continue

        # bullet
        m = re.match(r"^\s*[-*•]\s+(.+)$", line)
        if m:
            text = _strip_inline_md(m.group(1))
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(text)
            r.font.size = Pt(11)
            i += 1
            continue

        # 有序列表
        m = re.match(r"^\s*\d+[.)]\s+(.+)$", line)
        if m:
            text = _strip_inline_md(m.group(1))
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(text)
            r.font.size = Pt(11)
            i += 1
            continue

        # 引用
        if line.lstrip().startswith(">"):
            text = _strip_inline_md(line.lstrip().lstrip(">").strip())
            _add_para(doc, text, size=10, italic=True, color="555555", after=2)
            i += 1
            continue

        # 水平线
        if re.fullmatch(r"\s*-{3,}\s*", line) or re.fullmatch(r"\s*\*{3,}\s*", line):
            _add_para(doc, "─" * 30, size=9, color="888888",
                      align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
            i += 1
            continue

        # 默认普通段落，连续非空行合并
        buf = [line]
        j = i + 1
        while j < len(lines):
            nxt = lines[j].rstrip()
            if not nxt.strip():
                break
            if re.match(r"^(#{1,6})\s+", nxt):
                break
            if re.match(r"^\s*[-*•]\s+", nxt) or re.match(r"^\s*\d+[.)]\s+", nxt):
                break
            if nxt.lstrip().startswith(">"):
                break
            if nxt.lstrip().startswith("|"):
                break
            buf.append(nxt)
            j += 1
        para_text = _strip_inline_md(" ".join(buf))
        _add_para(doc, para_text, size=11, after=4)
        i = j

    _flush_table()


# ─────────────────────────── 主入口 ───────────────────────────

def generate_disclosure_docx(project, all_md_files: Iterable) -> bytes:
    """生成专利交底书 .docx Blob。

    Parameters
    ----------
    project : Project ORM
    all_md_files : Iterable[FileNode]  已按 name 排序的 markdown FileNode 列表
    """
    # 按章节 key 收集 md 内容
    sections: dict[str, list[str]] = {}
    md_list = list(all_md_files)
    for f in md_list:
        key = _classify_md(f.name or "")
        if not key:
            continue
        sections.setdefault(key, []).append(f.content or "")

    def _section(key: str) -> str:
        chunks = sections.get(key, [])
        return "\n\n".join(c for c in chunks if c and c.strip())

    # ───── 准备元数据 ─────
    domain_label = (
        project.custom_domain
        or DOMAIN_LABEL.get(project.domain, project.domain)
    )
    today = datetime.now().strftime("%Y-%m-%d")
    now_str = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    # 抽取 attachments：从 file tree 中找 root-user-{pid} 下的文件
    attachments = [
        f for f in md_list
        # 这里 md_list 只含 AI 输出；attachments 由调用方在 route 中传别名也行。
        # 为简洁，attachments 留给路由层稍后扩展；此处用空列表占位。
    ]
    # 路由层若想列 attachments，可在 project.files 里筛 root-user-{pid} 子节点
    user_attachments: list = getattr(project, "_user_attachments", []) or []

    # ───── 构建 docx ─────
    doc = Document()

    # 标题
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title_p.add_run("技术交底书")
    tr.bold = True
    tr.font.size = Pt(22)

    # 抬头表
    _add_para(doc, f"发明名称：{project.title}", size=13, bold=True, after=4)
    _add_para(doc, "申请类型：☑ 发明　　☐ 实用新型　　☐ 外观设计")
    _add_para(doc, f"技术领域：{domain_label}")
    _add_para(doc, f"本专利发明人：{PLACEHOLDER}")
    _add_para(doc, f"技术交底书撰写人：（员工）{project.owner_id}")
    _add_para(doc, f"技术问题联系人：{PLACEHOLDER}")
    _add_para(doc, f"联系人电话：{PLACEHOLDER}")
    _add_para(doc, f"联系人邮箱：{PLACEHOLDER}")
    _add_para(doc, f"填写日期：{today}")

    # 术语解释
    _add_para(doc, "术语解释：", size=11, bold=True, after=2)
    intake_notes = ""
    if project.intake_json and isinstance(project.intake_json, dict):
        intake_notes = (project.intake_json.get("notes") or "").strip()
    _add_para(doc, intake_notes or PLACEHOLDER, size=10, color="555555")

    # 项目简介（额外加一段，便于代理人快速了解）
    if project.description:
        _add_para(doc, "项目简介：", size=11, bold=True, after=2)
        _md_to_paragraphs(doc, project.description, base_heading_offset=3)

    # ───── 一、背景技术 ─────
    _add_heading(doc, "一、介绍背景技术，并描述已有的与本发明最相近似的现有技术方案", level=1)
    _add_meta(doc, "提示：可多写几个；该现有技术方案在公开出版物上有记载，最好提供出处或专利号。")
    bg = _section("background")
    if bg:
        _md_to_paragraphs(doc, bg, base_heading_offset=2)
    else:
        _add_para(doc, PLACEHOLDER)

    # ───── 二、现有技术缺点 ─────
    _add_heading(doc, "二、现有技术的缺点是什么（客观评价）？", level=1)
    _add_meta(doc, "提示：客观评价现有技术的缺点是针对本发明的优点来说的。可从结构、流程等角度推导（成本高、反应速度慢、结构复杂等）。")
    dis = _section("disadvantage")
    if dis:
        _md_to_paragraphs(doc, dis, base_heading_offset=2)
    else:
        _add_para(doc, PLACEHOLDER)

    # ───── 三、技术问题 ─────
    _add_heading(doc, "三、本发明解决的技术问题或技术目的？", level=1)
    _add_meta(doc, "提示：对应现有技术的所有缺点，正面描述本发明要解决的技术问题。")
    pb = _section("problem")
    if pb:
        _md_to_paragraphs(doc, pb, base_heading_offset=2)
    else:
        _add_para(doc, PLACEHOLDER)

    # ───── 四、技术方案 ─────
    _add_heading(doc, "四、本发明技术方案的详细阐述（即如何解决上述存在的问题）", level=1)
    _add_meta(doc, "提示：①核心是说明书公开的技术方案；②每一功能都要有相应的技术实现方案；③应当清楚、完整地描述技术特征及作用、原理。")
    sol = _section("solution")
    if sol:
        _md_to_paragraphs(doc, sol, base_heading_offset=2)
    else:
        _add_para(doc, PLACEHOLDER)

    # ───── 五、关键点和欲保护点 ─────
    _add_heading(doc, "五、本发明的关键点和欲保护点", level=1)
    _add_meta(doc, "提示：第四条是完整技术方案，本部分是提炼的关键创新点，列 1、2、3…，便于代理人撰写权利要求。")
    kp = _section("key_points")
    if kp:
        _md_to_paragraphs(doc, kp, base_heading_offset=2)
    else:
        _add_para(doc, PLACEHOLDER)

    # ───── 六、相比现有技术的优点 ─────
    _add_heading(doc, "六、与第一条所述的最好的现有技术相比，本发明有何优点？", level=1)
    _add_meta(doc, "提示：结合技术方案分条描述，以推理方式说明，做到有理有据；可对应第三条要解决的问题。")
    ad = _section("advantage")
    if ad:
        _md_to_paragraphs(doc, ad, base_heading_offset=2)
    else:
        _add_para(doc, PLACEHOLDER)

    # ───── 七、替代方案（从问题清单中抽取） ─────
    _add_heading(doc, "七、针对第四部分中的技术方案，是否还有别的替代方案同样能实现发明目的？", level=1)
    _add_meta(doc, "提示：拓展思路可以是部分结构、器件、方法步骤的替代，也可以是完整方案的替代。")
    questions_md = _section("questions")
    alt_text = _extract_alternatives(questions_md)
    if alt_text:
        _md_to_paragraphs(doc, alt_text, base_heading_offset=2)
    else:
        _add_para(doc, PLACEHOLDER)

    # ───── 八、是否经过实验 ─────
    _add_heading(doc, "八、本发明是否经过实验、模拟、使用而证明可行，结果如何？", level=1)
    _add_meta(doc, "提示：如果有，请简单说明并表述结果或效果即可。")
    exp_text = _extract_experiments(questions_md, _section("solution"))
    if exp_text:
        _md_to_paragraphs(doc, exp_text, base_heading_offset=2)
    else:
        _add_para(doc, PLACEHOLDER)

    # ───── 九、其他辅助资料 ─────
    _add_heading(doc, "九、其他有助于专利代理人理解本技术的资料", level=1)
    _add_meta(doc, "提示：技术文档、说明书等，可帮助代理人更快完成申请文件撰写。")
    if user_attachments:
        _add_para(doc, f"已上传 {len(user_attachments)} 项辅助资料：")
        for a in user_attachments:
            line = f"• {a.name}"
            if getattr(a, "url", None):
                line += f"  ({a.url})"
            elif getattr(a, "size", None):
                line += f"  ({a.size} bytes)"
            _add_para(doc, line, size=10, after=2)
    else:
        _add_para(doc, PLACEHOLDER)

    # 若有问题清单，附在最后供代理人参考
    if questions_md:
        _add_heading(doc, "附：AI 生成时遗留的问题清单", level=2)
        _md_to_paragraphs(doc, questions_md, base_heading_offset=3)

    # ───── 末尾 ─────
    doc.add_paragraph()
    _add_para(
        doc, "─" * 30, size=9, color="888888",
        align=WD_ALIGN_PARAGRAPH.CENTER, after=4,
    )
    _add_para(
        doc,
        f"本文档由 PatentlyPatent 自动生成 · 项目 {project.id} · {now_str}",
        size=9, color="888888", italic=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ─────────────────────────── heuristics ───────────────────────────

def _extract_alternatives(questions_md: str) -> str:
    """从 _问题清单.md 中抽取「替代方案」相关段落。"""
    if not questions_md:
        return ""
    lines = questions_md.splitlines()
    out: list[str] = []
    capture = False
    for ln in lines:
        if re.search(r"(替代方案|替代方式|alternative|可选方案|另一种)", ln, re.IGNORECASE):
            capture = True
            out.append(ln)
            continue
        if capture:
            if re.match(r"^#{1,6}\s+", ln):
                # 新章节，结束捕获
                if not re.search(r"(替代|alternative)", ln, re.IGNORECASE):
                    capture = False
                    continue
            out.append(ln)
            # 若连续 2 个空行，结束
            if len(out) >= 2 and not out[-1].strip() and not out[-2].strip():
                capture = False
    return "\n".join(out).strip()


def _extract_experiments(questions_md: str, solution_md: str) -> str:
    """从问题清单 / 技术方案中抽取实验/验证相关内容。"""
    chunks: list[str] = []
    for src in (questions_md, solution_md):
        if not src:
            continue
        for para in re.split(r"\n\s*\n", src):
            if re.search(r"(实验|测试|实测|验证|对比|结果|数据|benchmark|baseline)",
                         para, re.IGNORECASE):
                chunks.append(para.strip())
    return "\n\n".join(chunks).strip()

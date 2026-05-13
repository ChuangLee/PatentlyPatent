"""v0.37: 按 No.34 模板填充专利交底书 .docx

输入：project_id（从 DB 拉所有素材）
输出：bytes（docx 二进制）
模板：templates/专利技术交底书_No34.docx（一次性手动 .doc → .docx 入 git）

数据来源（按优先级、缺失时回退）：
  - "我的资料/0-报门.md"         → 标题 / 撰写人 / 联系信息（占位用 owner_id）
  - mineFull 5 节（.ai-internal/_compare/full/）：
      prior_art.md           → 章一+二（背景技术、现有技术缺点）
      summary.md             → 章三+六（技术问题、本发明优点）
      embodiments.md         → 章四（技术方案详细阐述）
      claims.md              → 章五（关键点和欲保护点）
      drawings_description.md → 章九补充材料
  - "AI 输出/_问题清单.md"        → 章七（替代方案）+ 章八（实验验证）
  - "AI 输出/调研下载/"           → 章九（其他资料：相似专利 + 相关文章列表）
  - "我的资料/" 用户上传的文件名  → 章九（参考资料）

填充策略（保模板版式）：
  - 不动模板原 9 章标题段
  - 不动模板斜体『使用说明』提示段（作为参考保留给员工）
  - 在每个提示段之后**插入填好的内容段**
  - 内容来源缺失 → 插入 "（请补充）" 占位，由用户在 Word 里手改

入口：
  - 库函数：generate_no34_docx(project_id, db_session=None) -> bytes
  - CLI：python -m app.disclosure_no34 <project_id> [-o out.docx]
"""
from __future__ import annotations

import io
import re
import sys
from copy import deepcopy
from datetime import datetime
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.document import Document as _Doc
from docx.oxml.ns import qn
from docx.shared import Pt

from .config import PROJECT_ROOT


TEMPLATE_PATH = PROJECT_ROOT / "templates" / "专利技术交底书_No34.docx"

# 模板第一段已经有"技术交底书"标题，第二段开始是占位字段；
# 第 10 段开始进入 9 章节（一/二/.../九）正文。
# 每章节结构：
#   段 X:   "一、xxxx"        ← chapter heading
#   段 X+1: 『...使用说明...』  ← italic helper（保留）
#   [段 X+2..]  ← 在这里插入填好的内容
#
# 我们用「章标题文本前缀」匹配段落定位，避免硬编码段号。

CHAPTER_PREFIXES = [
    ("一、", "background"),
    ("二、", "disadvantage"),
    ("三、", "problem"),
    ("四、", "solution"),
    ("五、", "key_points"),
    ("六、", "advantage"),
    ("七、", "alternative"),
    ("八、", "experiment"),
    ("九、", "other_material"),
]


# ─── Markdown → 干净文本 ──────────────────────────────────────────────────

_INJECT_RE = re.compile(r"<!--\s*\[LLM_INJECT::[^\]]+\]\s*-->")
_EXAMINER_RE = re.compile(r"<details>.*?</details>", re.DOTALL)
_HTML_TAG_RE = re.compile(r"</?(details|summary)[^>]*>", re.IGNORECASE)


def _strip_md(text: str) -> str:
    """去除 markdown 噪声留可读文本（粗体/code/链接/HTML 装饰）。"""
    if not text:
        return ""
    text = _INJECT_RE.sub("", text)
    text = _EXAMINER_RE.sub("", text)
    text = _HTML_TAG_RE.sub("", text)
    # ![alt](url) → alt
    text = re.sub(r"!\[([^\]]*)\]\([^)]+\)", r"\1", text)
    # [text](url) → text
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1", text)
    # `code` → code
    text = re.sub(r"`([^`]+)`", r"\1", text)
    # **bold** → bold
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"__([^_]+)__", r"\1", text)
    # *italic* → italic
    text = re.sub(r"(?<!\*)\*([^*\n]+)\*(?!\*)", r"\1", text)
    return text.strip()


def _md_lines_to_paragraphs(text: str) -> list[str]:
    """markdown 文本 → 段落列表（保留列表项缩进，丢弃前缀符号）。"""
    if not text:
        return []
    out: list[str] = []
    for line in text.splitlines():
        raw = line.rstrip()
        if not raw.strip():
            continue
        stripped = raw.lstrip()
        indent_lvl = (len(raw) - len(stripped)) // 2
        prefix = "  " * indent_lvl
        # 跳 markdown heading 自身（# / ## / ###），它已成章节
        if stripped.startswith("#"):
            txt = stripped.lstrip("#").strip()
            if txt:
                out.append(_strip_md(txt))
            continue
        # 列表项 - / * / 1.
        m = re.match(r"^([-*•]|\d+\.)\s+(.*)$", stripped)
        if m:
            out.append(f"{prefix}• {_strip_md(m.group(2))}")
            continue
        # 表格行 → 拼为 inline
        if stripped.startswith("|") and stripped.endswith("|"):
            cells = [c.strip() for c in stripped.strip("|").split("|") if c.strip()]
            if cells and not all(c.replace("-", "").replace(":", "") == "" for c in cells):
                out.append(_strip_md(" | ".join(cells)))
            continue
        # 引用 >
        if stripped.startswith(">"):
            out.append(_strip_md(stripped.lstrip(">").strip()))
            continue
        out.append(_strip_md(stripped))
    # 合并相邻空行
    return [p for p in out if p]


# ─── DB 抓素材 ───────────────────────────────────────────────────────────


def _gather_project_content(db, project_id: str) -> dict:
    """从 DB 抽所有 No34 需要的素材，返回 dict。"""
    from .models import FileNode, Project

    p = db.get(Project, project_id)
    if p is None:
        raise ValueError(f"project {project_id} not found")

    all_files = db.query(FileNode).filter(FileNode.project_id == project_id).all()
    by_id = {f.id: f for f in all_files}
    by_parent: dict[str | None, list[FileNode]] = {}
    for f in all_files:
        by_parent.setdefault(f.parent_id, []).append(f)

    def _find_folder(parent_id: str | None, name: str) -> FileNode | None:
        for f in by_parent.get(parent_id, []):
            if f.kind == "folder" and f.name == name:
                return f
        return None

    def _read_named(folder_id: str | None, name: str) -> str:
        for f in by_parent.get(folder_id, []):
            if f.kind == "file" and f.name == name:
                return f.content or ""
        return ""

    # 报门
    user_root = f"root-user-{project_id}"
    intake_md = _read_named(user_root, "0-报门.md")

    # 用户上传文件清单
    user_files: list[str] = []
    for f in by_parent.get(user_root, []):
        if f.kind == "file" and f.name != "0-报门.md":
            user_files.append(f.name)

    # mineFull 5 节（.ai-internal/_compare/full/）
    sections: dict[str, str] = {}
    internal = next(
        (f for f in by_parent.get(None, []) if f.kind == "folder" and f.name == ".ai-internal"),
        None,
    )
    if internal:
        compare = _find_folder(internal.id, "_compare")
        if compare:
            full = _find_folder(compare.id, "full")
            if full:
                for f in by_parent.get(full.id, []):
                    if f.kind == "file" and f.content:
                        sections[f.name] = f.content

    # AI 输出
    ai_root = next(
        (f for f in by_parent.get(None, []) if f.kind == "folder" and f.name == "AI 输出"),
        None,
    )
    legacy_sections: dict[str, str] = {}
    questions_md = ""
    research_files: list[str] = []
    if ai_root:
        for f in by_parent.get(ai_root.id, []):
            if f.kind == "file":
                if f.name == "_问题清单.md":
                    questions_md = f.content or ""
                elif f.name.endswith(".md"):
                    legacy_sections[f.name] = f.content or ""
        # 调研下载 子树
        download = _find_folder(ai_root.id, "调研下载")
        if download:
            for cat in by_parent.get(download.id, []):
                if cat.kind == "folder":
                    for item in by_parent.get(cat.id, []):
                        if item.kind == "file":
                            research_files.append(f"{cat.name}/{item.name}")

    return {
        "project": p,
        "intake_md": intake_md,
        "user_files": user_files,
        "sections": sections,          # mineFull 5 节
        "legacy_sections": legacy_sections,  # 老 mining 01-/02-/.../06-
        "questions_md": questions_md,
        "research_files": research_files,
    }


# ─── 章节内容生成 ────────────────────────────────────────────────────────


def _section_text(data: dict, key: str) -> str:
    """读取 mineFull 节，没有就回退老 mining 节。"""
    section_map = {
        "prior_art": ["prior_art.md", "01-背景技术.md"],
        "summary": ["summary.md", "02-现有技术缺点.md", "03-技术问题.md"],
        "embodiments": ["embodiments.md", "04-技术方案.md"],
        "claims": ["claims.md", "05-关键点.md"],
        "drawings": ["drawings_description.md"],
        "advantage": ["06-优点.md"],
    }
    names = section_map.get(key, [])
    for n in names:
        if n in data["sections"] and data["sections"][n].strip():
            return data["sections"][n]
        if n in data["legacy_sections"] and data["legacy_sections"][n].strip():
            return data["legacy_sections"][n]
    return ""


def _section_paragraphs(data: dict, key: str) -> list[str]:
    return _md_lines_to_paragraphs(_section_text(data, key))


def _build_chapter_content(data: dict, chapter_id: str) -> list[str]:
    """为每一章生成填充段落列表。空 → 返回 ['（请补充）']。"""
    out: list[str] = []
    if chapter_id == "background":
        # 一、背景技术 + 最相近现有技术
        prior = _section_paragraphs(data, "prior_art")
        if prior:
            out.extend(prior)
        # mineFull summary 第一段也可能含背景
    elif chapter_id == "disadvantage":
        # 二、现有技术缺点 —— summary 里通常有"现有技术缺点"小节，先扫
        sum_text = _section_text(data, "summary")
        # 尝试抽"缺点"段
        m = re.search(r"(现有技术缺点|缺点[：:]|## .*?(缺点|不足).*?)\n([\s\S]+?)(?=\n## |\Z)", sum_text)
        if m:
            out.extend(_md_lines_to_paragraphs(m.group(3)))
        else:
            # 老 mining 路径有专门 02-现有技术缺点.md
            txt = data["legacy_sections"].get("02-现有技术缺点.md", "")
            if txt:
                out.extend(_md_lines_to_paragraphs(txt))
    elif chapter_id == "problem":
        # 三、技术问题
        sum_text = _section_text(data, "summary")
        m = re.search(r"(技术问题|本发明.*解决.*问题|## .*?问题.*?)\n([\s\S]+?)(?=\n## |\Z)", sum_text)
        if m:
            out.extend(_md_lines_to_paragraphs(m.group(2)))
        else:
            out.extend(_md_lines_to_paragraphs(data["legacy_sections"].get("03-技术问题.md", "")))
    elif chapter_id == "solution":
        # 四、技术方案 = embodiments
        out.extend(_section_paragraphs(data, "embodiments"))
    elif chapter_id == "key_points":
        # 五、关键点 = claims
        out.extend(_section_paragraphs(data, "claims"))
    elif chapter_id == "advantage":
        # 六、优点
        sum_text = _section_text(data, "summary")
        m = re.search(r"(优点|有益效果|关键效果|## .*?效果.*?)\n([\s\S]+?)(?=\n## |\Z)", sum_text)
        if m:
            out.extend(_md_lines_to_paragraphs(m.group(2)))
        else:
            out.extend(_section_paragraphs(data, "advantage"))
    elif chapter_id == "alternative":
        # 七、替代方案 → _问题清单.md "替代方案"小节
        q = data["questions_md"]
        m = re.search(r"(替代方案|## .*?替代.*?)\n([\s\S]+?)(?=\n## |\Z)", q)
        if m:
            out.extend(_md_lines_to_paragraphs(m.group(2)))
    elif chapter_id == "experiment":
        # 八、实验验证 → _问题清单.md "实验"小节
        q = data["questions_md"]
        m = re.search(r"(实验|模拟|验证.*可行|## .*?实验.*?)\n([\s\S]+?)(?=\n## |\Z)", q)
        if m:
            out.extend(_md_lines_to_paragraphs(m.group(2)))
    elif chapter_id == "other_material":
        # 九、其他资料：用户上传文件清单 + 调研下载文献
        if data["user_files"]:
            out.append("申请人提供的资料：")
            for n in data["user_files"]:
                out.append(f"  • {n}")
        if data["research_files"]:
            out.append("调研期间发现的相关文献（AI 已落地到 AI 输出/调研下载/）：")
            for n in data["research_files"]:
                out.append(f"  • {n}")
        drawings = _section_paragraphs(data, "drawings")
        if drawings:
            out.append("")
            out.append("【附图说明】")
            out.extend(drawings)

    return out or ["（请补充）"]


# ─── 模板插入 ────────────────────────────────────────────────────────────


def _insert_para_after(paragraph, text: str):
    """在指定 paragraph 之后插入一个新段；返回新段。"""
    new_p_xml = deepcopy(paragraph._p)
    # 清空内容
    for child in list(new_p_xml):
        if child.tag.endswith("}r") or child.tag.endswith("}hyperlink"):
            new_p_xml.remove(child)
    paragraph._p.addnext(new_p_xml)
    # 用 docx 的 Paragraph 包装
    from docx.text.paragraph import Paragraph
    new_para = Paragraph(new_p_xml, paragraph._parent)
    new_para.text = text
    # 强制中文字体（保持模板观感）
    for run in new_para.runs:
        run.font.size = Pt(11)
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            from lxml.etree import SubElement
            rFonts = SubElement(rPr, qn("w:rFonts"))
        rFonts.set(qn("w:eastAsia"), "宋体")
    return new_para


def _is_chapter_heading(text: str, prefix: str) -> bool:
    """是否是某章标题段（以「一、」「二、」... 起头）。"""
    return text.startswith(prefix)


def _is_helper_para(text: str) -> bool:
    """是否模板里的『使用说明』提示段（『...』形式）。"""
    t = text.strip()
    return (t.startswith("『") and t.endswith("』")) or (
        t.startswith("『") and len(t) > 2
    )


def _fill_field_para(doc: _Doc, prefix: str, value: str) -> bool:
    """模板里 'X: '类字段段（如『发明名称：xxx』）填值。命中返回 True。"""
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix):
            # 整段重写为「前缀 + value」
            for run in list(p.runs):
                run.text = ""
            if p.runs:
                p.runs[0].text = f"{prefix}{value}"
            else:
                p.text = f"{prefix}{value}"
            return True
    return False


# ─── 入口 ────────────────────────────────────────────────────────────────


def generate_no34_docx(project_id: str, db=None) -> bytes:
    """生成 No34 模板填充后的 docx bytes。"""
    if db is None:
        from .db import SessionLocal
        db = SessionLocal()
        own_session = True
    else:
        own_session = False

    try:
        data = _gather_project_content(db, project_id)
    finally:
        if own_session:
            db.close()

    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError(f"No34 模板不存在：{TEMPLATE_PATH}")

    doc = Document(str(TEMPLATE_PATH))
    project = data["project"]
    intake = project.intake_json or {}

    # 1) 顶部字段填充（发明名称 / 撰写人 / 联系人）
    _fill_field_para(doc, "发明名称：", project.title or "")
    _fill_field_para(doc, "本专利发明人:", str(project.owner_id or ""))
    _fill_field_para(doc, "技术交底书撰写人:", str(project.owner_id or ""))
    # 联系信息：从 intake 拿，没有就留模板原样
    contact = (intake.get("notes") or "").strip()
    if contact and "@" in contact:
        _fill_field_para(doc, "联系人邮箱:", contact[:80])

    # 2) 章节定位：遍历所有段落，找章节标题，在斜体提示段后插入内容
    paragraphs = list(doc.paragraphs)
    chapter_indices: list[tuple[int, str]] = []  # (paragraph_index, chapter_id)
    for i, p in enumerate(paragraphs):
        text = (p.text or "").strip()
        for prefix, cid in CHAPTER_PREFIXES:
            if text.startswith(prefix):
                chapter_indices.append((i, cid))
                break

    # 从后往前插入，避免索引漂移
    for idx, chapter_id in reversed(chapter_indices):
        p = paragraphs[idx]
        # 找紧跟着的『...』使用说明段（可能 1 段或多段连续）
        anchor = p
        next_idx = idx + 1
        while next_idx < len(paragraphs):
            nt = (paragraphs[next_idx].text or "").strip()
            if _is_helper_para(nt):
                anchor = paragraphs[next_idx]
                next_idx += 1
                continue
            break
        # 在 anchor 之后插入内容段
        content_paras = _build_chapter_content(data, chapter_id)
        # 反向插入（每次插在 anchor 之后，相当于按正序排列）
        for line in reversed(content_paras):
            _insert_para_after(anchor, line)

    # 3) 输出
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ─── CLI ────────────────────────────────────────────────────────────────


def _cli() -> int:
    import argparse
    parser = argparse.ArgumentParser(description="按 No34 模板生成专利交底书 docx")
    parser.add_argument("project_id", help="项目 ID 如 p-c0b7566d")
    parser.add_argument("-o", "--output", help="输出路径（默认 ./<title>-交底书.docx）")
    args = parser.parse_args()

    blob = generate_no34_docx(args.project_id)

    out_path: Path
    if args.output:
        out_path = Path(args.output)
    else:
        # 默认从 DB 拿 title 取名
        from .db import SessionLocal
        from .models import Project
        with SessionLocal() as s:
            p = s.get(Project, args.project_id)
            title = (p.title if p else args.project_id) or args.project_id
        out_path = Path.cwd() / f"{title}-交底书.docx"

    out_path.write_bytes(blob)
    print(f"✓ 已生成 {out_path} ({len(blob)} bytes)")
    return 0


if __name__ == "__main__":
    sys.exit(_cli())
